from __future__ import annotations

import html
import json
import mimetypes
import os
import re
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

import chromadb
from sentence_transformers import SentenceTransformer

from db.database import Database
from db.models import Document, DocumentImportResult, RagContext, RagSearchResult


TOKEN_RE = re.compile(r"[a-zA-Zа-яА-ЯёЁ0-9]+")
TAG_RE = re.compile(r"<[^>]+>")
PROJECT_ROOT = Path(__file__).parent.parent


class RagTools:
    def __init__(
        self,
        db: Database,
        chunk_size: int = 900,
        chunk_overlap: int = 160,
        chroma_path: str | None = None,
        embedding_model_name: str | None = None,
    ):
        self.db = db
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.embedding_model_name = embedding_model_name or os.environ.get(
            "RAG_EMBEDDING_MODEL",
            "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
        )
        self.local_files_only = os.environ.get("RAG_LOCAL_FILES_ONLY", "0") == "1"
        self.chroma_path = chroma_path or os.environ.get(
            "CHROMA_PATH",
            str(PROJECT_ROOT / "chroma_db"),
        )
        self._embedding_model: SentenceTransformer | None = None
        self.chroma_client = chromadb.PersistentClient(path=self.chroma_path)
        self.collection = self.chroma_client.get_or_create_collection(
            name="university_documents",
            metadata={"hnsw:space": "cosine"},
        )

    def import_document(
        self,
        path: str,
        discipline_id: str | None = None,
        title: str | None = None,
    ) -> DocumentImportResult:
        source_path = Path(path).expanduser()
        if not source_path.is_absolute():
            source_path = Path.cwd() / source_path
        source_path = source_path.resolve()

        if not source_path.exists():
            raise FileNotFoundError(f"Document not found: {source_path}")
        if not source_path.is_file():
            raise ValueError(f"Document path is not a file: {source_path}")

        pages = self._extract_pages(source_path)
        chunks = self._chunk_pages(pages)
        if not chunks:
            raise ValueError(f"Document has no readable text: {source_path}")

        document_id = str(uuid.uuid4())
        mime_type = mimetypes.guess_type(source_path.name)[0] or "application/octet-stream"
        document_title = title or source_path.stem
        created_at = datetime.now(timezone.utc).isoformat()

        cursor = self.db.conn.cursor()
        existing = cursor.execute(
            "SELECT id FROM documents WHERE source_path = ?",
            (str(source_path),),
        ).fetchone()
        if existing:
            self._delete_document_vectors(existing["id"])
            cursor.execute("DELETE FROM document_chunks WHERE document_id = ?", (existing["id"],))
            cursor.execute("DELETE FROM documents WHERE id = ?", (existing["id"],))

        cursor.execute(
            """
            INSERT INTO documents (
                id, title, source_path, mime_type, discipline_id, created_at, metadata_json
            ) VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (
                document_id,
                document_title,
                str(source_path),
                mime_type,
                discipline_id,
                created_at,
                json.dumps(
                    {
                        "vector_store": "chromadb",
                        "embedding_model": self.embedding_model_name,
                        "chroma_collection": self.collection.name,
                    },
                    ensure_ascii=False,
                ),
            ),
        )

        chunk_ids: list[str] = []
        chunk_texts: list[str] = []
        chunk_metadatas: list[dict[str, str | int | float | bool]] = []

        for index, chunk in enumerate(chunks):
            content = chunk["content"]
            chunk_id = str(uuid.uuid4())
            cursor.execute(
                """
                INSERT INTO document_chunks (
                    id, document_id, chunk_index, page, content, embedding_json, token_count
                ) VALUES (?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    chunk_id,
                    document_id,
                    index,
                    chunk["page"],
                    content,
                    "[]",
                    len(self._tokenize(content)),
                ),
            )
            chunk_ids.append(chunk_id)
            chunk_texts.append(content)
            chunk_metadatas.append(
                {
                    "document_id": document_id,
                    "document_title": document_title,
                    "source_path": str(source_path),
                    "discipline_id": discipline_id or "",
                    "chunk_index": index,
                    "page": int(chunk["page"]) if chunk["page"] is not None else -1,
                }
            )

        self.collection.add(
            ids=chunk_ids,
            documents=chunk_texts,
            embeddings=self._embed_batch(chunk_texts),
            metadatas=chunk_metadatas,
        )
        self.db.conn.commit()

        document = Document(
            id=document_id,
            title=document_title,
            source_path=str(source_path),
            mime_type=mime_type,
            discipline_id=discipline_id,
            created_at=created_at,
        )
        return DocumentImportResult(document=document, chunks_count=len(chunks))

    def list_documents(self, discipline_id: str | None = None) -> list[Document]:
        cursor = self.db.conn.cursor()
        if discipline_id:
            rows = cursor.execute(
                """
                SELECT id, title, source_path, mime_type, discipline_id, created_at
                FROM documents
                WHERE discipline_id = ?
                ORDER BY created_at DESC
                """,
                (discipline_id,),
            ).fetchall()
        else:
            rows = cursor.execute(
                """
                SELECT id, title, source_path, mime_type, discipline_id, created_at
                FROM documents
                ORDER BY created_at DESC
                """
            ).fetchall()

        return [self._document_from_row(row) for row in rows]

    def search_documents(
        self,
        query: str,
        discipline_id: str | None = None,
        limit: int = 5,
    ) -> list[RagSearchResult]:
        normalized_query = query.strip()
        if not normalized_query:
            return []

        limit = max(1, min(limit, 20))
        where = {"discipline_id": discipline_id} if discipline_id else None
        query_result = self.collection.query(
            query_embeddings=self._embed_batch([normalized_query]),
            n_results=limit,
            where=where,
            include=["documents", "metadatas", "distances"],
        )

        ids = query_result.get("ids", [[]])[0]
        documents = query_result.get("documents", [[]])[0]
        metadatas = query_result.get("metadatas", [[]])[0]
        distances = query_result.get("distances", [[]])[0]

        results: list[RagSearchResult] = []
        for chunk_id, content, metadata, distance in zip(ids, documents, metadatas, distances):
            page = int(metadata["page"])
            discipline = str(metadata.get("discipline_id") or "")
            results.append(
                RagSearchResult(
                    document_id=str(metadata["document_id"]),
                    document_title=str(metadata["document_title"]),
                    source_path=str(metadata["source_path"]),
                    discipline_id=discipline or None,
                    chunk_id=str(chunk_id),
                    chunk_index=int(metadata["chunk_index"]),
                    page=page if page >= 0 else None,
                    score=round(max(0.0, 1.0 - float(distance)), 6),
                    content=str(content),
                )
            )

        return results

    def build_rag_context(
        self,
        query: str,
        discipline_id: str | None = None,
        limit: int = 5,
    ) -> RagContext:
        chunks = self.search_documents(query=query, discipline_id=discipline_id, limit=limit)
        return RagContext(
            query=query,
            answer_instruction=(
                "Ответь на вопрос только по найденным фрагментам документов. "
                "Если в контексте нет ответа, прямо скажи, что данных в документах недостаточно. "
                "Ссылайся на название документа и страницу, когда page заполнен."
            ),
            chunks=chunks,
        )

    def _extract_pages(self, source_path: Path) -> list[dict[str, int | str | None]]:
        suffix = source_path.suffix.lower()
        if suffix in {".txt", ".md", ".markdown", ".csv", ".json", ".py"}:
            return [{"page": None, "text": source_path.read_text(encoding="utf-8")}]
        if suffix in {".html", ".htm"}:
            raw = source_path.read_text(encoding="utf-8")
            text = html.unescape(TAG_RE.sub(" ", raw))
            return [{"page": None, "text": text}]
        if suffix == ".pdf":
            return self._extract_pdf_pages(source_path)
        if suffix == ".docx":
            return self._extract_docx_pages(source_path)
        raise ValueError(
            "Unsupported document format. Supported: txt, md, html, pdf, docx, csv, json, py."
        )

    def _extract_pdf_pages(self, source_path: Path) -> list[dict[str, int | str | None]]:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("PDF import requires optional dependency: pypdf") from exc

        reader = PdfReader(str(source_path))
        pages = []
        for index, page in enumerate(reader.pages, start=1):
            pages.append({"page": index, "text": page.extract_text() or ""})
        return pages

    def _extract_docx_pages(self, source_path: Path) -> list[dict[str, int | str | None]]:
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise RuntimeError("DOCX import requires optional dependency: python-docx") from exc

        document = DocxDocument(str(source_path))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        return [{"page": None, "text": text}]

    def _chunk_pages(self, pages: Iterable[dict[str, int | str | None]]) -> list[dict[str, int | str | None]]:
        chunks: list[dict[str, int | str | None]] = []
        for page in pages:
            text = self._normalize_text(str(page["text"] or ""))
            if not text:
                continue

            start = 0
            while start < len(text):
                end = min(start + self.chunk_size, len(text))
                if end < len(text):
                    boundary = text.rfind(" ", start, end)
                    if boundary > start + int(self.chunk_size * 0.65):
                        end = boundary

                content = text[start:end].strip()
                if content:
                    chunks.append({"page": page["page"], "content": content})

                if end >= len(text):
                    break
                start = max(0, end - self.chunk_overlap)
        return chunks

    def _embed_batch(self, texts: list[str]) -> list[list[float]]:
        embeddings = self.embedding_model.encode(
            texts,
            normalize_embeddings=True,
            show_progress_bar=False,
        )
        return embeddings.tolist()

    @property
    def embedding_model(self) -> SentenceTransformer:
        if self._embedding_model is None:
            try:
                self._embedding_model = SentenceTransformer(
                    self.embedding_model_name,
                    local_files_only=self.local_files_only,
                )
            except Exception as exc:
                raise RuntimeError(
                    "Failed to load RAG embedding model "
                    f"'{self.embedding_model_name}'. "
                    "Check internet access for the first download, or set "
                    "RAG_EMBEDDING_MODEL to a local model path. "
                    "If the model is already cached, you can set RAG_LOCAL_FILES_ONLY=1."
                ) from exc
        return self._embedding_model

    def _delete_document_vectors(self, document_id: str) -> None:
        self.collection.delete(where={"document_id": document_id})

    def _tokenize(self, text: str) -> list[str]:
        return [token.lower() for token in TOKEN_RE.findall(text)]

    def _normalize_text(self, text: str) -> str:
        lines = [line.strip() for line in text.replace("\x00", " ").splitlines()]
        return re.sub(r"\s+", " ", " ".join(line for line in lines if line)).strip()

    def _document_from_row(self, row) -> Document:
        return Document(
            id=row["id"],
            title=row["title"],
            source_path=row["source_path"],
            mime_type=row["mime_type"],
            discipline_id=row["discipline_id"],
            created_at=row["created_at"],
        )
