from __future__ import annotations

import base64
import json
import os
import random
import re
import textwrap
import urllib.error
import urllib.request
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable, Protocol, Sequence

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.shared import Pt
from faker import Faker

from agent_tutor_sdk.contracts import Discipline
from agent_tutor_sdk.data_client import DataServiceClientSync as DataServiceClient
from rag.fixtures._material import Material
from rag.fixtures.catalog import (
    CURRICULUM,
    DISCIPLINE_TOPICS,
    GROUP_NAMES,
    GROUP_SPECIALTY_MAP,
    SPECIALITIES,
)
from rag.fixtures.rag_tools import RagTools

# Выходная директория для сгенерированных материалов.
# По умолчанию — ./generated_materials относительно cwd.
# Переопределяется через env DOCGEN_OUTPUT_DIR.
DEFAULT_OUTPUT_DIR = Path(os.environ.get("DOCGEN_OUTPUT_DIR", "./generated_materials")).resolve()


DEFAULT_DOCGEN_MODEL = "qwen2.5:0.5b"
DOCX_FONT_NAME = "Times New Roman"
DOCX_BODY_SIZE_PT = 14

GENERIC_TOPICS = [
    "основные понятия дисциплины",
    "типовые задачи предметной области",
    "практические ограничения и допущения",
    "инструменты анализа результата",
    "ошибки, возникающие при применении методов",
]

DEPARTMENTS = [
    "кафедра прикладной информатики",
    "кафедра программной инженерии",
    "кафедра интеллектуальных систем",
    "кафедра информационной безопасности",
    "кафедра вычислительной математики",
]

GOAL_TEMPLATES = [
    "сформировать у студентов понимание ключевых идей дисциплины "
    "«{discipline}» и показать их применение в направлении "
    "«{specialty}».",
    "подготовить студентов к решению типовых учебных и прикладных задач "
    "по теме «{discipline}» с опорой на понятные алгоритмы действий.",
    "связать теоретические разделы дисциплины «{discipline}» с практикой "
    "проектирования, анализа и проверки решений.",
]

OUTCOME_TEMPLATES = [
    "объяснять тему «{topic}» простыми терминами и на учебных примерах",
    "выбирать подходящий метод решения задачи в рамках дисциплины",
    "проверять корректность результата и находить типовые ошибки",
    "оформлять решение в виде понятного отчета или краткого технического описания",
    "связывать материал курса с задачами направления «{specialty}»",
]

PRACTICE_CONTEXTS = [
    "учебной информационной системы кафедры",
    "сервиса учета лабораторных работ",
    "портала расписания и успеваемости",
    "небольшого аналитического модуля для преподавателя",
    "прототипа студенческого проекта",
]

CONTROL_FORMATS = [
    "мини-тест и устное обсуждение решения",
    "лабораторный отчет с демонстрацией результата",
    "практическая задача с коротким письменным объяснением",
    "разбор типовой ошибки и защита исправленного решения",
]

WEAK_RESPONSE_MARKERS = [
    "к сожалению",
    "не могу",
    "не могу предоставить",
    "как искусственный интеллект",
    "как языковая модель",
]


@dataclass(frozen=True)
class MaterialSpec:
    material_type: str
    extension: str
    title_prefix: str
    body_word_range: tuple[int, int] = (450, 800)

    @property
    def suffix(self) -> str:
        return f".{self.extension}"


@dataclass(frozen=True)
class GeneratedDocument:
    title: str
    material_type: str
    path: Path
    text: str


@dataclass(frozen=True)
class DocumentScenario:
    discipline: Discipline
    spec: MaterialSpec
    specialty: str
    group_name: str
    course: int
    semester: int
    teacher_name: str
    department: str
    goal: str
    topics: tuple[str, ...]
    outcomes: tuple[str, ...]
    practice_context: str
    control_format: str


@dataclass(frozen=True)
class OllamaGenerateConfig:
    model: str
    url: str
    timeout_seconds: float
    temperature: float
    num_predict: int
    min_response_chars: int
    max_attempts: int

    @classmethod
    def from_env(cls) -> OllamaGenerateConfig:
        return cls(
            model=os.environ.get("DOCGEN_MODEL", DEFAULT_DOCGEN_MODEL),
            url=os.environ.get("DOCGEN_OLLAMA_URL", _default_ollama_generate_url()),
            timeout_seconds=_env_float("DOCGEN_TIMEOUT", 3600.0),
            temperature=_env_float("DOCGEN_TEMPERATURE", 1.00),
            num_predict=_env_int("DOCGEN_NUM_PREDICT", 4500),
            min_response_chars=_env_int("DOCGEN_MIN_RESPONSE_CHARS", 120),
            max_attempts=max(1, _env_int("DOCGEN_MAX_ATTEMPTS", 2)),
        )


class TextGenerationClient(Protocol):
    def generate(
        self,
        prompt: str,
        images: Sequence[str | Path] | None = None,
    ) -> str: ...


class OllamaGenerateClient:
    """Small client for Ollama's /api/generate endpoint.

    Ollama accepts base64-encoded images for multimodal models. The document
    generator does not need images today, but keeping the hook here makes that a
    client concern instead of leaking HTTP payload details into generation code.
    """

    def __init__(self, config: OllamaGenerateConfig | None = None) -> None:
        self.config = config or OllamaGenerateConfig.from_env()

    def generate(
        self,
        prompt: str,
        images: Sequence[str | Path] | None = None,
    ) -> str:
        best_text = ""
        for attempt in range(self.config.max_attempts):
            current_prompt = prompt if attempt == 0 else self._retry_prompt(prompt)
            data = self._request(current_prompt, images)
            if error := data.get("error"):
                raise RuntimeError(f"Ollama returned an error: {error}")

            text = str(data.get("response", "")).strip()
            if len(text) > len(best_text):
                best_text = text
            if len(text) >= self.config.min_response_chars:
                return text

        if best_text:
            return best_text
        raise RuntimeError(f"Ollama model `{self.config.model}` returned empty text.")

    def _request(
        self,
        prompt: str,
        images: Sequence[str | Path] | None,
    ) -> dict[str, Any]:
        payload = self._payload(prompt, images)
        request = urllib.request.Request(
            self.config.url,
            data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
            headers={"Content-Type": "application/json"},
            method="POST",
        )

        try:
            with urllib.request.urlopen(
                request,
                timeout=self.config.timeout_seconds,
            ) as response:
                data = json.loads(response.read().decode("utf-8"))
        except urllib.error.HTTPError as exc:
            raise RuntimeError(self._http_error_message(exc)) from exc
        except (OSError, TimeoutError, json.JSONDecodeError) as exc:
            raise RuntimeError(
                "Ollama is unavailable. Start it with `ollama serve` and make "
                f"sure model `{self.config.model}` is installed."
            ) from exc

        return data

    def _payload(
        self,
        prompt: str,
        images: Sequence[str | Path] | None,
    ) -> dict[str, Any]:
        payload: dict[str, Any] = {
            "model": self.config.model,
            "prompt": prompt,
            "stream": False,
            "options": {
                "temperature": self.config.temperature,
                "num_predict": self.config.num_predict,
            },
        }
        if images:
            payload["images"] = [_image_to_base64(image) for image in images]
        return payload

    def _http_error_message(self, exc: urllib.error.HTTPError) -> str:
        details = exc.read().decode("utf-8", errors="replace").strip()
        if details:
            return f"Ollama request failed with HTTP {exc.code}: {details}"
        return f"Ollama request failed with HTTP {exc.code}."

    @staticmethod
    def _retry_prompt(prompt: str) -> str:
        return (
            f"{prompt}\n\n"
            "Предыдущий ответ был слишком коротким. Раскрой тему подробнее, "
            "добавь связные объяснения и 2-3 учебных примера."
        )


class DocumentScenarioFactory:
    """Builds repeatable-ish educational context around a discipline."""

    def __init__(self, seed: int | None = None) -> None:
        self.fake = Faker("ru_RU")
        self.random = random.Random(seed)
        if seed is not None:
            self.fake.seed_instance(seed)

    @classmethod
    def from_env(cls) -> DocumentScenarioFactory:
        return cls(seed=_env_optional_int("DOCGEN_FAKE_SEED"))

    def build(self, discipline: Discipline, spec: MaterialSpec) -> DocumentScenario:
        specialty = self._specialty_for_discipline(discipline.name)
        group_name = self._group_for_specialty(specialty)
        course = self.random.randint(1, 4)
        semester = course * 2 - self.random.randint(0, 1)
        topics = self._topics_for_discipline(discipline.name)

        return DocumentScenario(
            discipline=discipline,
            spec=spec,
            specialty=specialty,
            group_name=group_name,
            course=course,
            semester=semester,
            teacher_name=self.fake.name(),
            department=self.random.choice(DEPARTMENTS),
            goal=self._goal(discipline.name, specialty),
            topics=topics,
            outcomes=self._outcomes(topics, specialty),
            practice_context=self._practice_context(),
            control_format=self.random.choice(CONTROL_FORMATS),
        )

    def _specialty_for_discipline(self, discipline_name: str) -> str:
        matches = [
            specialty
            for specialty, discipline_names in CURRICULUM.items()
            if discipline_name in discipline_names
        ]
        return self.random.choice(matches or SPECIALITIES)

    def _group_for_specialty(self, specialty: str) -> str:
        matching_groups = []
        for group_name in GROUP_NAMES:
            prefix = group_name.split("-")[0]
            if GROUP_SPECIALTY_MAP.get(prefix) == specialty:
                matching_groups.append(group_name)
        return self.random.choice(matching_groups or GROUP_NAMES)

    def _topics_for_discipline(self, discipline_name: str) -> tuple[str, ...]:
        topics = DISCIPLINE_TOPICS.get(discipline_name, GENERIC_TOPICS)
        sample_size = min(len(topics), self.random.randint(4, 5))
        selected = set(self.random.sample(topics, sample_size))
        return tuple(topic for topic in topics if topic in selected)

    def _goal(self, discipline_name: str, specialty: str) -> str:
        template = self.random.choice(GOAL_TEMPLATES)
        return template.format(discipline=discipline_name, specialty=specialty)

    def _outcomes(self, topics: tuple[str, ...], specialty: str) -> tuple[str, ...]:
        templates = self.random.sample(OUTCOME_TEMPLATES, k=4)
        return tuple(
            template.format(
                topic=self.random.choice(topics),
                specialty=specialty,
            )
            for template in templates
        )

    def _practice_context(self) -> str:
        return f"{self.random.choice(PRACTICE_CONTEXTS)} «{self.fake.company()}»"


class DocumentTextGenerator:
    """Fills a Faker-built educational scaffold with LLM-generated body text."""

    def __init__(
        self,
        client: TextGenerationClient | None = None,
        scenario_factory: DocumentScenarioFactory | None = None,
    ) -> None:
        self.client = client or OllamaGenerateClient()
        self.scenario_factory = scenario_factory or DocumentScenarioFactory.from_env()

    def generate(
        self,
        discipline: Discipline,
        material: MaterialSpec | str,
    ) -> str:
        spec = (
            material if isinstance(material, MaterialSpec) else _material_spec(material)
        )
        scenario = self.scenario_factory.build(discipline, spec)
        body = self.client.generate(self.build_prompt(scenario))
        if _is_weak_body(body):
            body = _fallback_body(scenario)
        return self.render_document(scenario, body)

    @staticmethod
    def build_prompt(scenario: DocumentScenario) -> str:
        low, high = scenario.spec.body_word_range
        lines = [
            "Ты пишешь фрагмент учебного материала на русском языке.",
            "Паспорт, цель, план, вопросы и выводы будут добавлены автоматически.",
            "Верни связный учебный текст с Markdown-заголовками уровня ###.",
            f"Дисциплина: {scenario.discipline.name}",
            f"Описание дисциплины: {scenario.discipline.description}",
            f"Тип материала: {scenario.spec.material_type}",
            f"Направление: {scenario.specialty}",
            f"Курс: {scenario.course}, семестр: {scenario.semester}",
            f"Цель материала: {scenario.goal}",
            "Темы, которые нужно раскрыть:",
            *[f"- {topic}" for topic in scenario.topics],
            "Требования:",
            f"- ориентировочный объём: {low}-{high} слов;",
            "- используй Markdown-заголовки уровня ### для крупных фрагментов;",
            "- пиши связно, с короткими примерами и без выдуманных источников;",
            "- начни сразу с учебного объяснения.",
        ]
        return "\n".join(lines)

    @staticmethod
    def render_document(scenario: DocumentScenario, body: str) -> str:
        sections = [
            "## Паспорт материала",
            f"- Направление: {scenario.specialty}",
            f"- Группа: {scenario.group_name}",
            f"- Курс: {scenario.course}",
            f"- Семестр: {scenario.semester}",
            f"- Преподаватель: {scenario.teacher_name}",
            f"- Подразделение: {scenario.department}",
            "",
            "## Цель",
            scenario.goal,
            "",
            "## План",
            *_numbered_lines(scenario.topics),
            "",
            "## Ожидаемые результаты",
            *_bullet_lines(scenario.outcomes),
            "",
            "## Содержательная часть",
            body.strip(),
            "",
            "## Практический фрагмент",
            _practice_fragment(scenario),
            "",
            "## Контрольные вопросы",
            *_numbered_lines(_control_questions(scenario)),
            "",
            "## Формат проверки",
            scenario.control_format,
            "",
            "## Краткие выводы",
            _summary_fragment(scenario),
        ]
        return "\n".join(sections)


MATERIAL_SPECS: tuple[MaterialSpec, ...] = (
    MaterialSpec("Лекция", "pdf", "Лекция", (650, 1000)),
    MaterialSpec("Методичка", "docx", "Методичка", (500, 850)),
    MaterialSpec(
        "Лабораторная работа",
        "docx",
        "Лабораторная работа",
        (450, 750),
    ),
)


class MaterialDocumentGenerator:
    def __init__(
        self,
        rag_tools: RagTools,
        output_dir: str | Path | None = None,
        text_generator: DocumentTextGenerator | None = None,
        material_specs: Sequence[MaterialSpec] = MATERIAL_SPECS,
        data_client: "DataServiceClient | None" = None,
    ) -> None:
        self.rag_tools = rag_tools
        self._data_client = data_client or DataServiceClient()
        self.doc_repo = rag_tools.pipeline.repository
        self.output_dir = Path(
            output_dir or os.environ.get("DOCGEN_OUTPUT_DIR", DEFAULT_OUTPUT_DIR)
        )
        self.text_generator = text_generator or DocumentTextGenerator()
        self.material_specs = tuple(material_specs)

    def ensure_materials(
        self,
        discipline_id: str,
        force: bool = False,
    ) -> list[Material]:
        # Проверяем что дисциплина существует через data-service HTTP
        disciplines = self._data_client.get_all_disciplines()
        discipline = next((d for d in disciplines if d.id == discipline_id), None)
        if discipline is None:
            return []

        self._remove_stale_generated_documents(discipline_id, force=force)

        existing = self._valid_generated_materials(
            self.doc_repo.get_materials(discipline_id)
        )
        missing_types = self._expected_material_types() - {
            material.type for material in existing
        }
        if existing and not force and not missing_types:
            return existing

        generated = self.generate_documents(
            discipline,
            material_types=None if force else missing_types,
        )
        for document in generated:
            self._index_generated_document(discipline.id, document)

        return self.doc_repo.get_materials(discipline_id)

    def generate_documents(
        self,
        discipline: Discipline,
        material_types: set[str] | None = None,
    ) -> list[GeneratedDocument]:
        discipline_dir = self.output_dir / _slugify(discipline.name)
        discipline_dir.mkdir(parents=True, exist_ok=True)

        generated: list[GeneratedDocument] = []
        for spec in self.material_specs:
            if material_types is not None and spec.material_type not in material_types:
                continue
            generated.append(self._generate_document(discipline, spec, discipline_dir))
        return generated

    def _generate_document(
        self,
        discipline: Discipline,
        spec: MaterialSpec,
        discipline_dir: Path,
    ) -> GeneratedDocument:
        title = f"{spec.title_prefix}: {discipline.name}"
        text = self.text_generator.generate(discipline, spec)
        path = discipline_dir / self._file_name(spec, discipline)

        _write_document(path, title, text)
        return GeneratedDocument(
            title=title,
            material_type=spec.material_type,
            path=path,
            text=text,
        )

    def _file_name(self, spec: MaterialSpec, discipline: Discipline) -> str:
        prefix = _slugify(spec.title_prefix)
        discipline_name = _slugify(discipline.name)
        return f"{prefix}_{discipline_name}{spec.suffix}"

    def _index_generated_document(
        self,
        discipline_id: str,
        document: GeneratedDocument,
    ) -> None:
        try:
            chunks = self.rag_tools.pipeline.chunker.chunk_pages(
                [{"page": None, "text": document.text}]
            )
            self.rag_tools.pipeline.repository.save_document_with_chunks(
                source_path=str(document.path.resolve()),
                chunks=chunks,
                discipline_id=discipline_id,
                title=document.title,
                vector_store=self.rag_tools.pipeline.vector_store,
            )
        except Exception:
            self.rag_tools.pipeline.repository.save_generated_document_fallback(
                path=str(document.path),
                discipline_id=discipline_id,
                title=document.title,
                text=document.text,
            )

    def _remove_stale_generated_documents(
        self,
        discipline_id: str,
        force: bool,
    ) -> None:
        if force:
            self._delete_generated_documents_where(
                discipline_id=discipline_id,
                missing_only=False,
            )
            return

        self._delete_generated_documents_where(
            discipline_id=discipline_id,
            missing_only=True,
        )
        self._delete_outdated_generated_documents(discipline_id)

    def _delete_outdated_generated_documents(self, discipline_id: str) -> None:
        expected_extensions = self._expected_extensions()
        rows_to_delete = []
        for material in self.doc_repo.get_materials(discipline_id):
            source_path = Path(material.source_path)
            if not self._is_generated_path(source_path):
                continue

            expected_extension = expected_extensions.get(material.type)
            if expected_extension and source_path.suffix.lower() != expected_extension:
                rows_to_delete.append(
                    {"id": material.id, "source_path": material.source_path}
                )

        self._delete_document_rows(rows_to_delete)

    def _delete_generated_documents_where(
        self,
        discipline_id: str,
        missing_only: bool,
    ) -> None:
        rows = self.doc_repo.list_generated_document_rows(
            path_marker=self.output_dir.name,
            discipline_id=discipline_id,
        )

        rows_to_delete = []
        for row in rows:
            source_path = Path(row["source_path"])
            if not self._is_generated_path(source_path):
                continue
            if missing_only and source_path.exists():
                continue
            rows_to_delete.append(row)

        self._delete_document_rows(rows_to_delete)

    def _delete_document_rows(self, rows: Iterable[Any]) -> None:
        deleted_any = False
        for row in rows:
            self._delete_generated_document_row(row)
            deleted_any = True

        if deleted_any:
            self._cleanup_empty_output_dirs()

    def _delete_generated_document_row(self, row: Any) -> None:
        document_id = row["id"]
        source_path = Path(row["source_path"])

        try:
            self.rag_tools._delete_document_vectors(document_id)
        except Exception:
            pass

        self.doc_repo.delete_document_record(document_id, commit=False)
        if source_path.exists():
            try:
                source_path.unlink()
            except OSError:
                pass

    def _cleanup_empty_output_dirs(self) -> None:
        if not self.output_dir.exists():
            return

        paths = sorted(
            self.output_dir.rglob("*"),
            key=lambda item: len(item.parts),
            reverse=True,
        )
        for path in paths:
            if not path.is_dir():
                continue
            try:
                path.rmdir()
            except OSError:
                pass

        try:
            self.output_dir.rmdir()
        except OSError:
            pass

    def _valid_generated_materials(self, materials: list[Material]) -> list[Material]:
        expected_extensions = self._expected_extensions()
        valid = []
        for material in materials:
            source_path = Path(material.source_path)
            expected_extension = expected_extensions.get(material.type)
            if (
                expected_extension
                and self._is_generated_path(source_path)
                and source_path.suffix.lower() != expected_extension
            ):
                continue
            valid.append(material)
        return valid

    def _expected_extensions(self) -> dict[str, str]:
        return {spec.material_type: spec.suffix for spec in self.material_specs}

    def _expected_material_types(self) -> set[str]:
        return {spec.material_type for spec in self.material_specs}

    def _is_generated_path(self, source_path: Path) -> bool:
        if self.output_dir.name not in source_path.parts:
            return False

        try:
            source_path.resolve().relative_to(self.output_dir.resolve())
            return True
        except (OSError, ValueError):
            return False


def _material_spec(material_type: str) -> MaterialSpec:
    for spec in MATERIAL_SPECS:
        if spec.material_type == material_type:
            return spec
    return MaterialSpec(material_type, "docx", material_type)


def _numbered_lines(items: Sequence[str]) -> list[str]:
    return [f"{index}. {item}" for index, item in enumerate(items, 1)]


def _bullet_lines(items: Sequence[str]) -> list[str]:
    return [f"- {item}" for item in items]


def _practice_fragment(scenario: DocumentScenario) -> str:
    first_topic = scenario.topics[0]
    last_topic = scenario.topics[-1]
    return (
        f"Рассмотрите фрагмент {scenario.practice_context}. Необходимо описать, "
        f"как тема «{first_topic}» влияет на выбор решения, и показать, какие "
        f"ограничения появляются при работе с темой «{last_topic}». Результат "
        "оформляется как короткий отчёт: постановка задачи, ход рассуждения, "
        "полученный вывод и одно возможное улучшение."
    )


def _control_questions(scenario: DocumentScenario) -> tuple[str, ...]:
    questions = [
        f"Какую роль играет тема «{topic}» в дисциплине «{scenario.discipline.name}»?"
        for topic in scenario.topics[:3]
    ]
    questions.append(
        f"Какие ошибки вероятны при применении материала в направлении "
        f"«{scenario.specialty}»?"
    )
    questions.append(
        "Как проверить, что предложенное решение действительно соответствует "
        "поставленной задаче?"
    )
    return tuple(questions)


def _summary_fragment(scenario: DocumentScenario) -> str:
    topics = ", ".join(scenario.topics[:3])
    return (
        f"Материал фиксирует базовый маршрут по дисциплине "
        f"«{scenario.discipline.name}»: {topics}. Основной акцент сделан на "
        "связи понятий с практической задачей, проверке результата и аккуратном "
        "оформлении решения."
    )


def _is_weak_body(text: str) -> bool:
    normalized = text.strip().lower()
    if len(normalized) < 120:
        return True
    return any(marker in normalized for marker in WEAK_RESPONSE_MARKERS)


def _fallback_body(scenario: DocumentScenario) -> str:
    paragraphs = []
    for topic in scenario.topics:
        paragraphs.extend(
            [
                f"### {topic.capitalize()}",
                (
                    f"Тема «{topic}» рассматривается как часть дисциплины "
                    f"«{scenario.discipline.name}» и помогает студентам "
                    "перейти от общего определения к практическому действию. "
                    f"В направлении «{scenario.specialty}» этот раздел важен "
                    "потому, что он задает язык описания задачи и критерии "
                    "проверки результата."
                ),
                (
                    f"На учебном примере {scenario.practice_context} студент "
                    "может выделить исходные данные, выбрать способ работы с "
                    "ними, описать ожидаемый результат и проверить, где решение "
                    "может дать сбой. Такой формат хорошо подходит для "
                    "лабораторной работы, короткого отчета или обсуждения на "
                    "практическом занятии."
                ),
            ]
        )
    return "\n\n".join(paragraphs)


def _default_ollama_generate_url() -> str:
    host = os.environ.get("OLLAMA_HOST", "http://127.0.0.1:11434").rstrip("/")
    if not host.startswith(("http://", "https://")):
        host = f"http://{host}"
    return f"{host}/api/generate"


def _env_int(name: str, default: int) -> int:
    value = os.environ.get(name)
    if value is None:
        return default
    try:
        return int(value)
    except ValueError:
        return default


def _env_float(name: str, default: float) -> float:
    value = os.environ.get(name)
    if value is None:
        return default
    try:
        return float(value)
    except ValueError:
        return default


def _env_optional_int(name: str) -> int | None:
    value = os.environ.get(name)
    if value is None:
        return None
    try:
        return int(value)
    except ValueError:
        return None


def _image_to_base64(image: str | Path) -> str:
    image_path = Path(image)
    try:
        exists = image_path.exists()
    except OSError:
        exists = False
    if exists:
        return base64.b64encode(image_path.read_bytes()).decode("ascii")
    return str(image)


def _slugify(value: str) -> str:
    value = value.lower().replace("ё", "е")
    value = "".join(TRANSLIT.get(char, char) for char in value)
    value = re.sub(r"[^a-z0-9_]+", "_", value)
    value = re.sub(r"_+", "_", value).strip("_")
    return value or "material"


TRANSLIT = {
    "а": "a",
    "б": "b",
    "в": "v",
    "г": "g",
    "д": "d",
    "е": "e",
    "ж": "zh",
    "з": "z",
    "и": "i",
    "й": "y",
    "к": "k",
    "л": "l",
    "м": "m",
    "н": "n",
    "о": "o",
    "п": "p",
    "р": "r",
    "с": "s",
    "т": "t",
    "у": "u",
    "ф": "f",
    "х": "h",
    "ц": "ts",
    "ч": "ch",
    "ш": "sh",
    "щ": "sch",
    "ъ": "",
    "ы": "y",
    "ь": "",
    "э": "e",
    "ю": "yu",
    "я": "ya",
    " ": "_",
}


def _write_document(path: Path, title: str, text: str) -> None:
    if path.suffix == ".docx":
        _write_docx(path, title, text)
        return
    if path.suffix == ".pdf":
        _write_pdf(path, title, text)
        return
    raise ValueError(f"Unsupported generated document extension: {path.suffix}")


def _write_docx(path: Path, title: str, text: str) -> None:
    document = DocxDocument()
    _configure_docx_styles(document)

    title_paragraph = document.add_heading("", 0)
    _add_markdown_runs(title_paragraph, title)

    for block in _markdown_blocks(text):
        _add_docx_block(document, block)

    document.save(str(path))


def _add_docx_block(document: Any, block: str) -> None:
    heading_match = re.match(r"^(#{1,4})\s+(.+)$", block)
    if heading_match:
        level = min(len(heading_match.group(1)), 3)
        paragraph = document.add_heading("", level=level)
        _add_markdown_runs(paragraph, heading_match.group(2))
        return

    if re.match(r"^[-*]\s+", block):
        paragraph = document.add_paragraph(style="List Bullet")
        _add_markdown_runs(paragraph, re.sub(r"^[-*]\s+", "", block))
        return

    if re.match(r"^\d+[.)]\s+", block):
        paragraph = document.add_paragraph(style="List Number")
        _add_markdown_runs(paragraph, re.sub(r"^\d+[.)]\s+", "", block))
        return

    paragraph = document.add_paragraph()
    _add_markdown_runs(paragraph, block)


def _configure_docx_styles(document: Any) -> None:
    for style_name in ["Normal", "List Bullet", "List Number"]:
        _configure_docx_font(document.styles[style_name], DOCX_BODY_SIZE_PT)

    for style_name, size, bold in [
        ("Title", 18, True),
        ("Heading 1", 16, True),
        ("Heading 2", 15, True),
        ("Heading 3", 14, True),
    ]:
        _configure_docx_font(document.styles[style_name], size, bold=bold)


def _configure_docx_font(style: Any, size: int, bold: bool = False) -> None:
    style.font.name = DOCX_FONT_NAME
    style.font.size = Pt(size)
    style.font.bold = bold
    style._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_FONT_NAME)


def _add_markdown_runs(paragraph: Any, text: str) -> None:
    for part, is_bold in _split_bold(text):
        run = paragraph.add_run(_strip_inline_markdown(part))
        run.bold = is_bold
        run.font.name = DOCX_FONT_NAME
        run.font.size = Pt(DOCX_BODY_SIZE_PT)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), DOCX_FONT_NAME)


def _write_pdf(path: Path, title: str, text: str) -> None:
    wrapped_lines = [title, ""]
    for paragraph in _markdown_blocks(text):
        clean_paragraph = _strip_markdown(paragraph)
        wrapped_lines.extend(textwrap.wrap(clean_paragraph, width=74) or [""])
        wrapped_lines.append("")

    pages = list(_paginate(wrapped_lines, lines_per_page=36))
    objects = _build_pdf_objects(pages)
    _write_pdf_objects(path, objects)


def _build_pdf_objects(pages: list[list[str]]) -> list[bytes]:
    objects: list[bytes] = []
    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")

    page_refs = " ".join(f"{3 + index * 2} 0 R" for index in range(len(pages)))
    objects.append(
        f"<< /Type /Pages /Kids [{page_refs}] /Count {len(pages)} >>".encode()
    )

    font_object_id = 3 + len(pages) * 2
    for index, page_lines in enumerate(pages):
        page_id = 3 + index * 2
        content_id = page_id + 1
        objects.append(_pdf_page_object(page_id, content_id, font_object_id))
        objects.append(_pdf_stream_object(_pdf_text_stream(page_lines)))

    objects.extend(_pdf_font_objects(font_object_id))
    return objects


def _pdf_page_object(page_id: int, content_id: int, font_object_id: int) -> bytes:
    return (
        f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] "
        f"/Resources << /Font << /F1 {font_object_id} 0 R >> >> "
        f"/Contents {content_id} 0 R >>"
    ).encode()


def _pdf_stream_object(stream: bytes) -> bytes:
    return (
        b"<< /Length "
        + str(len(stream)).encode()
        + b" >>\nstream\n"
        + stream
        + b"\nendstream"
    )


def _pdf_font_objects(font_object_id: int) -> list[bytes]:
    cmap = _to_unicode_cmap()
    return [
        (
            b"<< /Type /Font /Subtype /Type0 /BaseFont /Times-Roman "
            b"/Encoding /Identity-H /DescendantFonts ["
            + str(font_object_id + 1).encode()
            + b" 0 R] /ToUnicode "
            + str(font_object_id + 2).encode()
            + b" 0 R >>"
        ),
        (
            b"<< /Type /Font /Subtype /CIDFontType2 /BaseFont /Times-Roman "
            b"/CIDSystemInfo << /Registry (Adobe) /Ordering (Identity) "
            b"/Supplement 0 >> "
            b"/FontDescriptor "
            + str(font_object_id + 3).encode()
            + b" 0 R /W [0 [600]] >>"
        ),
        _pdf_stream_object(cmap),
        (
            b"<< /Type /FontDescriptor /FontName /Times-Roman /Flags 4 "
            b"/FontBBox [-166 -225 1000 931] /ItalicAngle 0 /Ascent 931 "
            b"/Descent -225 /CapHeight 718 /StemV 80 >>"
        ),
    ]


def _markdown_blocks(text: str) -> list[str]:
    blocks: list[str] = []
    buffer: list[str] = []

    def flush_buffer() -> None:
        if buffer:
            blocks.append(" ".join(buffer).strip())
            buffer.clear()

    for raw_line in text.replace("\r\n", "\n").split("\n"):
        line = raw_line.strip()
        if not line:
            flush_buffer()
            continue
        if re.fullmatch(r"#{1,6}", line):
            flush_buffer()
            continue
        if re.match(r"^(#{1,6}\s*|[-*]\s+|\d+[.)]\s+)", line):
            flush_buffer()
            blocks.append(line)
            continue
        buffer.append(line)

    flush_buffer()
    return blocks


def _split_bold(text: str) -> list[tuple[str, bool]]:
    parts: list[tuple[str, bool]] = []
    position = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        if match.start() > position:
            parts.append((text[position : match.start()], False))
        parts.append((match.group(1), True))
        position = match.end()
    if position < len(text):
        parts.append((text[position:], False))
    return parts or [(text, False)]


def _strip_inline_markdown(text: str) -> str:
    return re.sub(r"`([^`]+)`", r"\1", text).replace("***", "").replace("__", "")


def _strip_markdown(text: str) -> str:
    text = re.sub(r"^#{1,6}\s*", "", text)
    text = re.sub(r"^[-*]\s+", "- ", text)
    text = re.sub(
        r"^\d+[.)]\s+",
        lambda match: match.group(0).replace(")", "."),
        text,
    )
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    return _strip_inline_markdown(text)


def _paginate(lines: list[str], lines_per_page: int) -> Iterable[list[str]]:
    for index in range(0, len(lines), lines_per_page):
        yield lines[index : index + lines_per_page]


def _pdf_text_stream(lines: list[str]) -> bytes:
    stream = ["BT", "/F1 14 Tf", "50 790 Td", "18 TL"]
    for line in lines:
        stream.append(f"<{line.encode('utf-16-be').hex()}> Tj")
        stream.append("T*")
    stream.append("ET")
    return "\n".join(stream).encode("ascii")


def _to_unicode_cmap() -> bytes:
    return (
        "/CIDInit /ProcSet findresource begin\n"
        "12 dict begin\n"
        "begincmap\n"
        "/CIDSystemInfo << /Registry (Adobe) /Ordering (UCS) /Supplement 0 >> def\n"
        "/CMapName /Adobe-Identity-UCS def\n"
        "/CMapType 2 def\n"
        "1 begincodespacerange\n"
        "<0000> <FFFF>\n"
        "endcodespacerange\n"
        "1 beginbfrange\n"
        "<0000> <FFFF> <0000>\n"
        "endbfrange\n"
        "endcmap\n"
        "CMapName currentdict /CMap defineresource pop\n"
        "end\n"
        "end"
    ).encode("ascii")


def _write_pdf_objects(path: Path, objects: list[bytes]) -> None:
    content = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for index, obj in enumerate(objects, 1):
        offsets.append(len(content))
        content.extend(f"{index} 0 obj\n".encode("ascii"))
        content.extend(obj)
        content.extend(b"\nendobj\n")

    xref_offset = len(content)
    content.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    content.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        content.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    content.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    path.write_bytes(bytes(content))
