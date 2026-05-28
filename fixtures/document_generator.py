from __future__ import annotations

import json
import os
import re
import textwrap
import urllib.error
import urllib.request
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.shared import Pt

from db.database import Database, PROJECT_ROOT
from db.models import Discipline, Material
from tools.rag import RagTools


MATERIAL_SPECS = [
    ("Лекция", "pdf", "Лекция"),
    ("Методичка", "docx", "Методичка"),
    ("Лабораторная работа", "docx", "Лабораторная работа"),
]


@dataclass(frozen=True)
class GeneratedDocument:
    title: str
    material_type: str
    path: Path
    text: str


class DocumentTextGenerator:
    """Generates educational text through a local Ollama endpoint."""

    def __init__(self) -> None:
        self.model = os.environ.get("DOCGEN_MODEL", "qwen2.5:0.5b")
        self.ollama_url = os.environ.get(
            "DOCGEN_OLLAMA_URL",
            self._default_ollama_url(),
        )

    @staticmethod
    def _default_ollama_url() -> str:
        host = os.environ.get("OLLAMA_HOST", "http://127.0.0.1:11434").rstrip("/")
        if not host.startswith(("http://", "https://")):
            host = f"http://{host}"
        return f"{host}/api/generate"

    def generate(self, discipline: Discipline, material_type: str) -> str:
        prompt = self._prompt(discipline, material_type)
        return self._from_ollama(prompt)

    @staticmethod
    def _prompt(discipline: Discipline, material_type: str) -> str:
        return (
            "Сгенерируй учебный материал для университета на русском языке.\n"
            f"Дисциплина: {discipline.name}\n"
            f"Описание дисциплины: {discipline.description}\n"
            f"Тип материала: {material_type}\n"
            "Требования:\n"
            "- объём 1200-1800 слов;\n"
            "- структурированный учебный текст с разделами, подпунктами и примерами;\n"
            "- используй Markdown только для заголовков (#, ##, ###), списков и **жирного текста**;\n"
            "- добавь практический пример, контрольные вопросы и краткие выводы;\n"
            "- без выдуманных ссылок и без обращения к читателю от лица ассистента."
        )

    def _from_ollama(self, prompt: str) -> str:
        payload = json.dumps(
            {
                "model": self.model,
                "prompt": prompt,
                "stream": False,
                "options": {
                    "temperature": 0.35,
                    "num_predict": int(os.environ.get("DOCGEN_NUM_PREDICT", "4500")),
                },
            },
            ensure_ascii=False,
        ).encode("utf-8")
        request = urllib.request.Request(
            self.ollama_url,
            data=payload,
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        try:
            with urllib.request.urlopen(request, timeout=90) as response:
                data = json.loads(response.read().decode("utf-8"))
        except (urllib.error.URLError, TimeoutError) as exc:
            raise RuntimeError(
                "Ollama is unavailable. Start it with `ollama serve` and make "
                f"sure model `{self.model}` is installed."
            ) from exc

        text = str(data.get("response", "")).strip()
        if len(text) < 300:
            raise RuntimeError(
                f"Ollama model `{self.model}` returned too little text."
            )
        return text


class MaterialDocumentGenerator:
    def __init__(
        self,
        db: Database,
        rag_tools: RagTools,
        output_dir: str | Path | None = None,
    ) -> None:
        self.db = db
        self.rag_tools = rag_tools
        self.output_dir = Path(
            output_dir
            or os.environ.get("DOCGEN_OUTPUT_DIR", PROJECT_ROOT / "generated_materials")
        )
        self.text_generator = DocumentTextGenerator()

    def ensure_materials(
        self,
        discipline_id: str,
        force: bool = False,
    ) -> list[Material]:
        discipline = self.db.get_discipline(discipline_id)
        if discipline is None:
            return []

        if force:
            self._delete_generated_documents(discipline_id)
        else:
            self._delete_missing_generated_documents(discipline_id)
            self._delete_outdated_generated_documents(discipline_id)

        existing = self.db.get_materials(discipline_id)
        existing = self._filter_valid_generated_materials(existing)
        expected_types = {spec[0] for spec in MATERIAL_SPECS}
        existing_types = {material.type for material in existing}
        missing_types = expected_types - existing_types
        if existing and not force and not missing_types:
            return existing

        generated = self.generate_documents(
            discipline,
            material_types=None if force else missing_types,
        )
        for document in generated:
            self._index_generated_document(discipline.id, document)

        return self.db.get_materials(discipline_id)

    def _delete_generated_documents(self, discipline_id: str) -> None:
        self._delete_generated_documents_where(
            discipline_id=discipline_id,
            missing_only=False,
        )

    def _delete_missing_generated_documents(self, discipline_id: str) -> None:
        self._delete_generated_documents_where(
            discipline_id=discipline_id,
            missing_only=True,
        )

    def _delete_outdated_generated_documents(self, discipline_id: str) -> None:
        expected_extensions = {material_type: extension for material_type, extension, _ in MATERIAL_SPECS}
        outdated_ids = []
        for material in self.db.get_materials(discipline_id):
            source_path = Path(material.source_path)
            if self.output_dir.name not in source_path.parts:
                continue
            expected_extension = expected_extensions.get(material.type)
            if expected_extension and source_path.suffix.lower() != f".{expected_extension}":
                outdated_ids.append(material.id)

        if not outdated_ids:
            return

        cursor = self.db.conn.cursor()
        for document_id in outdated_ids:
            row = cursor.execute(
                "SELECT id, source_path FROM documents WHERE id = ?",
                (document_id,),
            ).fetchone()
            if not row:
                continue
            self._delete_generated_document_row(cursor, row)
        self.db.conn.commit()

    def _delete_generated_documents_where(
        self,
        discipline_id: str,
        missing_only: bool,
    ) -> None:
        cursor = self.db.conn.cursor()
        rows = cursor.execute(
            """
            SELECT id, source_path FROM documents
            WHERE discipline_id = ? AND source_path LIKE ?
            """,
            (discipline_id, f"%{self.output_dir.name}%"),
        ).fetchall()

        for row in rows:
            if missing_only and Path(row["source_path"]).exists():
                continue
            self._delete_generated_document_row(cursor, row)
        self.db.conn.commit()
        self._cleanup_empty_output_dirs()

    def _delete_generated_document_row(self, cursor, row) -> None:
        source_path = Path(row["source_path"])
        try:
            self.rag_tools._delete_document_vectors(row["id"])
        except Exception:
            pass
        cursor.execute("DELETE FROM document_chunks WHERE document_id = ?", (row["id"],))
        cursor.execute("DELETE FROM documents WHERE id = ?", (row["id"],))
        if source_path.exists():
            try:
                source_path.unlink()
            except OSError:
                pass

    def _cleanup_empty_output_dirs(self) -> None:
        if not self.output_dir.exists():
            return
        for path in sorted(self.output_dir.rglob("*"), key=lambda item: len(item.parts), reverse=True):
            if not path.is_dir():
                continue
            try:
                path.rmdir()
            except OSError:
                pass
        try:
            self.output_dir.rmdir()
        except OSError:
            pass

    def _filter_valid_generated_materials(self, materials: list[Material]) -> list[Material]:
        expected_extensions = {material_type: extension for material_type, extension, _ in MATERIAL_SPECS}
        valid = []
        for material in materials:
            source_path = Path(material.source_path)
            expected_extension = expected_extensions.get(material.type)
            if (
                expected_extension
                and self.output_dir.name in source_path.parts
                and source_path.suffix.lower() != f".{expected_extension}"
            ):
                continue
            valid.append(material)
        return valid

    def generate_documents(
        self,
        discipline: Discipline,
        material_types: set[str] | None = None,
    ) -> list[GeneratedDocument]:
        discipline_dir = self.output_dir / _slugify(discipline.name)
        discipline_dir.mkdir(parents=True, exist_ok=True)

        generated: list[GeneratedDocument] = []
        for material_type, extension, title_prefix in MATERIAL_SPECS:
            if material_types is not None and material_type not in material_types:
                continue
            title = f"{title_prefix}: {discipline.name}"
            text = self.text_generator.generate(discipline, material_type)
            file_name = f"{_slugify(title_prefix)}_{_slugify(discipline.name)}.{extension}"
            path = discipline_dir / file_name
            if extension == "docx":
                _write_docx(path, title, text)
            else:
                _write_pdf(path, title, text)
            generated.append(
                GeneratedDocument(
                    title=title,
                    material_type=material_type,
                    path=path,
                    text=text,
                )
            )
        return generated

    def _index_generated_document(
        self,
        discipline_id: str,
        document: GeneratedDocument,
    ) -> None:
        try:
            chunks = self.rag_tools.pipeline.chunker.chunk_pages(
                [{"page": None, "text": document.text}]
            )
            self.rag_tools.pipeline._save_document(
                source_path=document.path.resolve(),
                chunks=chunks,
                discipline_id=discipline_id,
                title=document.title,
            )
        except Exception:
            self.db.save_generated_document_record(
                path=str(document.path),
                discipline_id=discipline_id,
                title=document.title,
                text=document.text,
            )

def _slugify(value: str) -> str:
    value = value.lower().replace("ё", "е")
    translit = {
        "а": "a", "б": "b", "в": "v", "г": "g", "д": "d",
        "е": "e", "ж": "zh", "з": "z", "и": "i", "й": "y",
        "к": "k", "л": "l", "м": "m", "н": "n", "о": "o",
        "п": "p", "р": "r", "с": "s", "т": "t", "у": "u",
        "ф": "f", "х": "h", "ц": "ts", "ч": "ch", "ш": "sh",
        "щ": "sch", "ъ": "", "ы": "y", "ь": "", "э": "e",
        "ю": "yu", "я": "ya", " ": "_",
    }
    value = "".join(translit.get(char, char) for char in value)
    value = re.sub(r"[^a-z0-9_]+", "_", value)
    value = re.sub(r"_+", "_", value).strip("_")
    return value or "material"


def _write_docx(path: Path, title: str, text: str) -> None:
    document = DocxDocument()
    _configure_docx_styles(document)

    title_paragraph = document.add_heading("", 0)
    _add_markdown_runs(title_paragraph, title)

    for block in _markdown_blocks(text):
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", block)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            paragraph = document.add_heading("", level=level)
            _add_markdown_runs(paragraph, heading_match.group(2))
        elif re.match(r"^[-*]\s+", block):
            paragraph = document.add_paragraph(style="List Bullet")
            _add_markdown_runs(paragraph, re.sub(r"^[-*]\s+", "", block))
        elif re.match(r"^\d+[.)]\s+", block):
            paragraph = document.add_paragraph(style="List Number")
            _add_markdown_runs(paragraph, re.sub(r"^\d+[.)]\s+", "", block))
        else:
            paragraph = document.add_paragraph()
            _add_markdown_runs(paragraph, block)
    document.save(path)


def _write_pdf(path: Path, title: str, text: str) -> None:
    wrapped_lines = [title, ""]
    for paragraph in _markdown_blocks(text):
        clean_paragraph = _strip_markdown(paragraph)
        wrapped_lines.extend(textwrap.wrap(clean_paragraph, width=74) or [""])
        wrapped_lines.append("")

    pages = list(_paginate(wrapped_lines, lines_per_page=36))
    objects: list[bytes] = []

    objects.append(b"<< /Type /Catalog /Pages 2 0 R >>")
    page_refs = " ".join(f"{3 + i * 2} 0 R" for i in range(len(pages)))
    objects.append(f"<< /Type /Pages /Kids [{page_refs}] /Count {len(pages)} >>".encode())

    font_object_id = 3 + len(pages) * 2
    for index, page_lines in enumerate(pages):
        page_id = 3 + index * 2
        content_id = page_id + 1
        objects.append(
            (
                f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] "
                f"/Resources << /Font << /F1 {font_object_id} 0 R >> >> "
                f"/Contents {content_id} 0 R >>"
            ).encode()
        )
        stream = _pdf_text_stream(page_lines)
        objects.append(
            b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n"
            + stream
            + b"\nendstream"
        )

    objects.append(
        b"<< /Type /Font /Subtype /Type0 /BaseFont /Times-Roman "
        b"/Encoding /Identity-H /DescendantFonts ["
        + str(font_object_id + 1).encode()
        + b" 0 R] /ToUnicode "
        + str(font_object_id + 2).encode()
        + b" 0 R >>"
    )
    objects.append(
        b"<< /Type /Font /Subtype /CIDFontType2 /BaseFont /Times-Roman "
        b"/CIDSystemInfo << /Registry (Adobe) /Ordering (Identity) /Supplement 0 >> "
        b"/FontDescriptor "
        + str(font_object_id + 3).encode()
        + b" 0 R /W [0 [600]] >>"
    )
    cmap = _to_unicode_cmap()
    objects.append(
        b"<< /Length " + str(len(cmap)).encode() + b" >>\nstream\n"
        + cmap
        + b"\nendstream"
    )
    objects.append(
        b"<< /Type /FontDescriptor /FontName /Times-Roman /Flags 4 "
        b"/FontBBox [-166 -225 1000 931] /ItalicAngle 0 /Ascent 931 "
        b"/Descent -225 /CapHeight 718 /StemV 80 >>"
    )

    _write_pdf_objects(path, objects)


def _configure_docx_styles(document: DocxDocument) -> None:
    styles = document.styles
    for style_name in ["Normal", "List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(14)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for style_name, size, bold in [
        ("Title", 18, True),
        ("Heading 1", 16, True),
        ("Heading 2", 15, True),
        ("Heading 3", 14, True),
    ]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _markdown_blocks(text: str) -> list[str]:
    blocks: list[str] = []
    buffer: list[str] = []

    def flush_buffer() -> None:
        if buffer:
            blocks.append(" ".join(buffer).strip())
            buffer.clear()

    for raw_line in text.replace("\r\n", "\n").split("\n"):
        line = raw_line.strip()
        if not line:
            flush_buffer()
            continue
        if re.fullmatch(r"#{1,6}", line):
            flush_buffer()
            continue
        if re.match(r"^(#{1,6}\s*|[-*]\s+|\d+[.)]\s+)", line):
            flush_buffer()
            blocks.append(line)
            continue
        buffer.append(line)

    flush_buffer()
    return blocks


def _add_markdown_runs(paragraph, text: str) -> None:
    for part, is_bold in _split_bold(text):
        run = paragraph.add_run(_strip_inline_markdown(part))
        run.bold = is_bold
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def _split_bold(text: str) -> list[tuple[str, bool]]:
    parts: list[tuple[str, bool]] = []
    position = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        if match.start() > position:
            parts.append((text[position:match.start()], False))
        parts.append((match.group(1), True))
        position = match.end()
    if position < len(text):
        parts.append((text[position:], False))
    return parts or [(text, False)]


def _strip_inline_markdown(text: str) -> str:
    return re.sub(r"`([^`]+)`", r"\1", text).replace("***", "").replace("__", "")


def _strip_markdown(text: str) -> str:
    text = re.sub(r"^#{1,6}\s*", "", text)
    text = re.sub(r"^[-*]\s+", "- ", text)
    text = re.sub(r"^\d+[.)]\s+", lambda match: match.group(0).replace(")", "."), text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    return _strip_inline_markdown(text)


def _paginate(lines: list[str], lines_per_page: int) -> Iterable[list[str]]:
    for index in range(0, len(lines), lines_per_page):
        yield lines[index:index + lines_per_page]


def _pdf_text_stream(lines: list[str]) -> bytes:
    stream = ["BT", "/F1 14 Tf", "50 790 Td", "18 TL"]
    for line in lines:
        stream.append(f"<{line.encode('utf-16-be').hex()}> Tj")
        stream.append("T*")
    stream.append("ET")
    return "\n".join(stream).encode("ascii")


def _to_unicode_cmap() -> bytes:
    return (
        "/CIDInit /ProcSet findresource begin\n"
        "12 dict begin\n"
        "begincmap\n"
        "/CIDSystemInfo << /Registry (Adobe) /Ordering (UCS) /Supplement 0 >> def\n"
        "/CMapName /Adobe-Identity-UCS def\n"
        "/CMapType 2 def\n"
        "1 begincodespacerange\n"
        "<0000> <FFFF>\n"
        "endcodespacerange\n"
        "1 beginbfrange\n"
        "<0000> <FFFF> <0000>\n"
        "endbfrange\n"
        "endcmap\n"
        "CMapName currentdict /CMap defineresource pop\n"
        "end\n"
        "end"
    ).encode("ascii")


def _write_pdf_objects(path: Path, objects: list[bytes]) -> None:
    content = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for index, obj in enumerate(objects, 1):
        offsets.append(len(content))
        content.extend(f"{index} 0 obj\n".encode("ascii"))
        content.extend(obj)
        content.extend(b"\nendobj\n")
    xref_offset = len(content)
    content.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    content.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        content.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    content.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    path.write_bytes(bytes(content))
